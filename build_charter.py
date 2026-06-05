"""
Generates two documents:
  1. Charter_ABP_EMEA_updated.docx  – concise one-page excerpt (portrait A4)
  2. Regulation_Special_Membership.docx – detailed rules referenced by the charter
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def set_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    s = doc.sections[0]
    s.top_margin    = Cm(top)
    s.bottom_margin = Cm(bottom)
    s.left_margin   = Cm(left)
    s.right_margin  = Cm(right)


def tight(p, before=0, after=3, line=13):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    fmt.line_spacing = Pt(line)
    return p


def keep_together(p):
    """Prevent a paragraph from being split across pages."""
    pPr = p._p.get_or_add_pPr()
    kl = OxmlElement('w:keepLines')
    pPr.append(kl)
    kn = OxmlElement('w:keepNext')
    pPr.append(kn)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11) if level == 1 else Pt(10.5)
    tight(p, before=4, after=2, line=13)
    return p


def add_body(doc, text, indent=False, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.size = Pt(size)
    tight(p, before=0, after=2, line=12)
    return p


def add_clause(doc, num, text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"{num} ")
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    tight(p, before=0, after=2, line=12)
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(size)
    tight(p, before=0, after=1, line=12)
    return p


def add_inline_clause(doc, label, text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    tight(p, before=0, after=2, line=12)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — CHARTER EXCERPT (concise, ~1 page)
# ══════════════════════════════════════════════════════════════════════════════
charter = Document()
set_margins(charter, top=2.0, bottom=2.0, left=2.5, right=2.5)

# Title
title = charter.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title.add_run("EXCERPT FROM THE CHARTER\nAssociation of Business Professionals (EMEA)")
rt.bold = True
rt.font.size = Pt(12)
tight(title, before=0, after=6, line=14)

# ── Section 1 ─────────────────────────────────────────────────────────────────
add_heading(charter, "1. Mission, Purpose, and Objectives")

add_heading(charter, "1.1. Mission", level=2)
add_body(charter,
    "The mission of the Association is to establish, unite, and support a professional community "
    "of highly qualified and internationally recognized individuals whose activities contribute to "
    "the advancement of entrepreneurship, strategic management, international trade, digital "
    "transformation, business innovation, and sustainable economic development within the EMEA "
    "region and globally.")

add_heading(charter, "1.2. Purpose", level=2)
add_body(charter,
    "The Association is established as a non-profit professional organization aimed at fostering "
    "collaboration, knowledge exchange, professional excellence, and the promotion of internationally "
    "recognized standards in business leadership and innovation.")

add_heading(charter, "1.3. Objectives", level=2)
add_body(charter, "To achieve its mission, the Association shall pursue the following objectives:")

for num, text in [
    ("1.3.1.", "To identify and bring together leading professionals in the fields of strategic management, international trade, supply chain management, product development, corporate governance, and business innovation whose activities demonstrate significant professional impact."),
    ("1.3.2.", "To develop and maintain an international professional network comprising executives, entrepreneurs, innovators, consultants, educators, and other industry leaders engaged in the advancement of sustainable business practices and modern economic models."),
    ("1.3.3.", "To promote the professional achievements and contributions of its Members at international level and to participate in global discussions related to digital transformation, technology-driven products, platform-based economies, creative industries, and innovation ecosystems."),
    ("1.3.4.", "To support continuous professional development, encourage adherence to high ethical and quality standards, and foster responsible and sustainable approaches to digital product development, financial technologies, corporate innovation, and user-centered design methodologies."),
    ("1.3.5.", "To contribute to the development and dissemination of best practices in business management, international cooperation, and innovation leadership."),
]:
    add_clause(charter, num, text)

# ── Section 2 ─────────────────────────────────────────────────────────────────
add_heading(charter, "2. Membership")

add_heading(charter, "2.1. Eligibility", level=2)
add_body(charter,
    "Membership in the Association shall be open to individuals who meet the professional and "
    "reputational standards established by the Association and who satisfy the criteria defined in "
    "this Charter and in the internal regulations adopted by the Association.")

add_heading(charter, "2.2. Membership Criteria", level=2)
add_body(charter, "Candidates for membership must demonstrate:")

for num, text in [
    ("2.2.1.", "A high level of professional expertise, leadership experience, strategic competence, and recognized contributions in the fields of business, international trade, supply chain management, entrepreneurship, product development, or related disciplines."),
    ("2.2.2.", "Evidence of international professional recognition, including but not limited to: leadership roles, significant industry achievements, participation in international projects, advisory activities, or comparable contributions to the development of business practices."),
    ("2.2.3.", "Documented professional engagement, which may include publications, academic or industry articles, media appearances, public speaking engagements, teaching activities, conference participation, advisory services, or other forms of recognized professional contribution."),
    ("2.2.4.", "Demonstrable impact on the development of business ecosystems, digital products, innovation initiatives, strategic transformation programs, or related areas within the EMEA region and internationally."),
]:
    add_clause(charter, num, text)

add_heading(charter, "2.3. Admission Procedure", level=2)
add_body(charter,
    "The procedure for admission, evaluation of candidates, approval of membership, suspension, and "
    "termination of membership shall be governed by internal regulations approved by the competent "
    "governing body of the Association.")

add_heading(charter, "2.4. Principles of Membership", level=2)
add_body(charter, "The Association shall operate on the basis of the principles of:")
for item in [
    "equality of Members;",
    "transparency in governance;",
    "professional integrity;",
    "ethical conduct;",
    "non-discrimination;",
    "voluntary participation;",
    "compliance with applicable laws and regulations.",
]:
    add_bullet(charter, item)

# ── 2.5 Honorary (Senior) Members — CONCISE ──────────────────────────────────
add_heading(charter, "2.5. Honorary (Senior) Members", level=2)
add_body(charter,
    "The Association may grant Honorary (Senior) Member status to individuals who have demonstrated "
    "outstanding achievements in the field of business. Candidates are accepted on the basis of "
    "criteria established for outstanding professional contribution to business development, "
    "as further specified in the Regulation on Special Membership approved by the governing body "
    "of the Association.")
add_body(charter,
    "Honorary (Senior) Members enjoy the rights and privileges of membership and are exempt from "
    "membership fees. The conditions of admission, criteria, rights, and revocation of this status "
    "are set out in full in the Regulation on Special Membership.")

# ── 2.6 Admissions Committee — CONCISE ───────────────────────────────────────
add_heading(charter, "2.6. Admissions Committee", level=2)
add_body(charter,
    "The Association shall maintain a standing Admissions Committee responsible for evaluating "
    "applications and nominations for all categories of membership, including Honorary (Senior) "
    "Members. The composition, functions, meeting procedure, minutes requirements, and "
    "confidentiality rules of the Admissions Committee are governed by the Regulation on Special "
    "Membership.")

# ── Certification ─────────────────────────────────────────────────────────────
# Force the certification block to stay together on the same page
cert_head = charter.add_paragraph()
cert_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = cert_head.add_run("Certification")
r.bold = True
r.font.size = Pt(10.5)
tight(cert_head, before=6, after=2, line=13)
keep_together(cert_head)

cert_text = charter.add_paragraph()
cert_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = cert_text.add_run("This excerpt is issued in accordance with the original Charter of the Association.")
r.font.size = Pt(10)
tight(cert_text, before=0, after=4, line=12)
keep_together(cert_text)

# Signature image
sig_para = charter.add_paragraph()
sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
sig_para.add_run().add_picture("/workspace/Подпись устав .png", width=Inches(1.8))
tight(sig_para, before=0, after=2, line=12)
keep_together(sig_para)

# Date below signature
date_para = charter.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
rd = date_para.add_run("Date: November 17, 2025")
rd.font.size = Pt(10)
tight(date_para, before=0, after=4, line=12)
keep_together(date_para)

# Name block
name_para = charter.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
rn = name_para.add_run("Domenico Liguori")
rn.bold = True
rn.font.size = Pt(10)
tight(name_para, before=0, after=1, line=12)
keep_together(name_para)

for line in ["President", "Association of Business Professionals (EMEA)", "info@emeaabp.org"]:
    p = charter.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(line)
    r.font.size = Pt(10)
    tight(p, before=0, after=1, line=12)
    keep_together(p)

charter.save("/workspace/Charter_ABP_EMEA_updated.docx")
print("Saved: Charter_ABP_EMEA_updated.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 — REGULATION ON SPECIAL MEMBERSHIP (detailed)
# ══════════════════════════════════════════════════════════════════════════════
reg = Document()
set_margins(reg, top=2.5, bottom=2.5, left=3.0, right=2.0)

# Title
title2 = reg.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt2 = title2.add_run(
    "REGULATION ON SPECIAL MEMBERSHIP\n"
    "Association of Business Professionals (EMEA)"
)
rt2.bold = True
rt2.font.size = Pt(13)
tight(title2, before=0, after=8, line=16)

sub = reg.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Adopted pursuant to Section 2.5 and Section 2.6 of the Charter")
rs.italic = True
rs.font.size = Pt(10)
tight(sub, before=0, after=10, line=13)

# ── Article 1 — Honorary (Senior) Members ────────────────────────────────────
add_heading(reg, "Article 1. Honorary (Senior) Members")

add_heading(reg, "1.1. Definition", level=2)
add_body(reg,
    "An Honorary (Senior) Member is an individual recognized by the Association for exceptional and "
    "sustained contributions to the field of business, including but not limited to: founding or "
    "scaling transformative enterprises, shaping industry standards, advancing international business "
    "cooperation, or mentoring the next generation of business leaders. Honorary (Senior) Member "
    "status is a distinction of honor and does not carry the obligations of regular membership, "
    "unless the individual voluntarily assumes them.", size=10)

add_heading(reg, "1.2. Eligibility Criteria", level=2)
add_body(reg,
    "A candidate for Honorary (Senior) Member status must satisfy at least three (3) of the "
    "following criteria:", size=10)

for item in [
    "(a) has held a senior executive, board, or advisory position in one or more organizations of international significance for a cumulative period of not less than ten (10) years;",
    "(b) has made a measurable and recognized contribution to the development of a business sector, market, or professional community at the regional (EMEA) or global level, evidenced by industry reports, independent assessments, peer recognition, or comparable sources;",
    "(c) has been the recipient of internationally recognized professional honors, awards, fellowships, or distinctions in the field of business, entrepreneurship, innovation, or economic development;",
    "(d) has authored, co-authored, or substantially contributed to publications, research, frameworks, or standards that have demonstrably influenced business practices, governance models, or industry thought leadership;",
    "(e) has served as a mentor, educator, or advocate for emerging business leaders, contributing to the professional development of others in a sustained and verifiable manner;",
    "(f) has represented the business community in international forums, governmental advisory bodies, intergovernmental organizations, or multi-stakeholder platforms in a capacity that advanced the interests of the business profession.",
]:
    add_bullet(reg, item)

add_heading(reg, "1.3. Nomination and Admission Procedure", level=2)
add_body(reg,
    "Candidates for Honorary (Senior) Member status may be nominated by: (i) any two (2) current "
    "Members of the Association; (ii) the President of the Association; or (iii) the Admissions "
    "Committee acting on its own initiative. Nominations must be submitted in writing and accompanied "
    "by a statement of justification addressing the eligibility criteria set out in Article 1.2. "
    "The Admissions Committee shall review the nomination and make a written recommendation to the "
    "governing body of the Association. The final decision to confer Honorary (Senior) Member status "
    "shall be adopted by the governing body by a simple majority of votes.", size=10)

add_heading(reg, "1.4. Rights and Privileges", level=2)
add_body(reg,
    "Honorary (Senior) Members shall be entitled to: use the designation 'Honorary Member' or "
    "'Senior Member' of the Association; participate in Association events, conferences, and "
    "professional activities; receive communications and publications issued by the Association; "
    "and be listed in the official register of Honorary Members maintained by the Association. "
    "Honorary (Senior) Members are exempt from membership fees unless they elect to contribute voluntarily.", size=10)

add_heading(reg, "1.5. Revocation", level=2)
add_body(reg,
    "Honorary (Senior) Member status may be revoked by the governing body of the Association in the "
    "event of conduct that is materially inconsistent with the values, principles, or reputation of "
    "the Association, following a review by the Admissions Committee.", size=10)

# ── Article 2 — Admissions Committee ─────────────────────────────────────────
add_heading(reg, "Article 2. Admissions Committee")

add_heading(reg, "2.1. Composition", level=2)
add_body(reg,
    "The Admissions Committee shall consist of not fewer than three (3) and not more than seven (7) "
    "members appointed by the governing body of the Association for a term of two (2) years, "
    "renewable once. The Committee shall be chaired by a Chairperson elected from among its members. "
    "At least one member of the Committee shall hold Honorary (Senior) Member status, where such "
    "members exist. Committee members must be in good standing with the Association and must not "
    "have a conflict of interest in respect of any application or nomination under review.", size=10)

add_heading(reg, "2.2. Functions and Powers", level=2)
add_body(reg, "The Admissions Committee shall:", size=10)
for item in [
    "(a) receive and process applications and nominations for membership;",
    "(b) verify that candidates satisfy the applicable eligibility criteria;",
    "(c) conduct due diligence on submitted documentation;",
    "(d) issue written recommendations to the governing body with respect to the admission, deferral, or rejection of each candidate;",
    "(e) maintain a confidential register of all applications reviewed;",
    "(f) conduct periodic reviews of membership criteria and recommend updates to the governing body;",
    "(g) initiate or conduct reviews in cases of potential revocation of membership or Honorary (Senior) Member status.",
]:
    add_bullet(reg, item)

add_heading(reg, "2.3. Procedure of Meetings", level=2)
add_body(reg,
    "The Admissions Committee shall meet as required and not less than four (4) times per year, "
    "whether in person or by electronic means. Meetings shall be convened by the Chairperson or, "
    "in their absence, by any two (2) Committee members. A quorum shall consist of a majority of "
    "the Committee members then in office. Decisions shall be adopted by a simple majority of "
    "members present and voting; in the event of a tie, the Chairperson shall have a casting vote.", size=10)

add_heading(reg, "2.4. Minutes of Meetings", level=2)
add_body(reg,
    "Minutes shall be drawn up for each meeting of the Admissions Committee. The minutes shall record:",
    size=10)
for item in [
    "the date, time, and location (or virtual platform) of the meeting;",
    "the names of members present and absent;",
    "the agenda items discussed;",
    "a summary of the deliberations;",
    "the decisions adopted, including the outcome of each application or nomination reviewed (approved, deferred, or rejected) with brief reasons;",
    "any dissenting opinions where a member so requests.",
]:
    add_bullet(reg, item)
add_body(reg,
    "Minutes shall be signed by the Chairperson and the designated Secretary of the meeting, "
    "approved at the following meeting, and kept in the official records of the Association for a "
    "minimum of five (5) years. Copies of minutes shall be made available to the governing body "
    "upon request. Personal data contained in the minutes shall be processed in accordance with "
    "applicable data protection laws.", size=10)

add_heading(reg, "2.5. Confidentiality", level=2)
add_body(reg,
    "All deliberations of the Admissions Committee and the content of individual applications and "
    "nominations shall be treated as strictly confidential. Committee members shall not disclose "
    "information relating to specific candidates outside the Committee, except as required by the "
    "governing body of the Association or by applicable law.", size=10)

# ── Certification ─────────────────────────────────────────────────────────────
reg.add_paragraph()
cert2 = reg.add_paragraph()
cert2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = cert2.add_run("Certification")
r.bold = True
r.font.size = Pt(10.5)
tight(cert2, before=6, after=2, line=13)
keep_together(cert2)

ct2 = reg.add_paragraph()
ct2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = ct2.add_run(
    "This Regulation is adopted in accordance with the Charter of the Association and shall "
    "enter into force upon approval by the governing body."
)
r.font.size = Pt(10)
tight(ct2, before=0, after=4, line=12)
keep_together(ct2)

sig2 = reg.add_paragraph()
sig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
sig2.add_run().add_picture("/workspace/Подпись устав .png", width=Inches(1.8))
tight(sig2, before=0, after=2, line=12)
keep_together(sig2)

date2 = reg.add_paragraph()
date2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = date2.add_run("Date: November 17, 2025")
r.font.size = Pt(10)
tight(date2, before=0, after=4, line=12)
keep_together(date2)

name2 = reg.add_paragraph()
name2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = name2.add_run("Domenico Liguori")
r.bold = True
r.font.size = Pt(10)
tight(name2, before=0, after=1, line=12)
keep_together(name2)

for line in ["President", "Association of Business Professionals (EMEA)", "info@emeaabp.org"]:
    p = reg.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(line)
    r.font.size = Pt(10)
    tight(p, before=0, after=1, line=12)
    keep_together(p)

reg.save("/workspace/Regulation_Special_Membership.docx")
print("Saved: Regulation_Special_Membership.docx")
