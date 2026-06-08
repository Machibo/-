from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    if level == 1:
        hs.font.size = Pt(20)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(15)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FA"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('ПОРТРЕТ ПОКУПАТЕЛЯ', bold=True, size=28,
         color=(0x1B, 0x3A, 0x5C), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('ДЛЯ НАСТРОЙКИ GOOGLE ADS', bold=True, size=20,
         color=(0x1B, 0x3A, 0x5C), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para('MYRONYX Global Awards  •  Global 1000 Award  •  NOVA Awards  •  Glonary Awards  •  Velto European Awards',
         size=12, color=(0x55, 0x55, 0x55), align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_para('Июнь 2026', size=12, color=(0x55, 0x55, 0x55), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Конфиденциально', size=11, italic=True, color=(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Содержание', level=1)
toc_items = [
    'Часть 1. Кто этот человек',
    'Часть 2. Что он хочет получить',
    'Часть 3. Как он ведёт себя в интернете',
    'Часть 4. Психология — почему он купит',
    'Часть 5. Настройки Google Ads — конкретные параметры',
    'Часть 6. Примеры объявлений',
    'Часть 7. Путь клиента от рекламы до оплаты',
    'Часть 8. Метрики и KPI',
    'Часть 9. Резюме — кто этот американец',
]
for item in toc_items:
    add_para(item, size=12, color=(0x1B, 0x3A, 0x5C))

doc.add_page_break()

# ============================================================
# PART 1
# ============================================================
doc.add_heading('Часть 1. Кто этот человек', level=1)

doc.add_heading('1.1. Общее описание', level=2)
add_para(
    'Это американский предприниматель или руководитель компании, мужчина или женщина 33–48 лет, '
    'который владеет или управляет бизнесом с выручкой от $500 тыс. до $10 млн в год. '
    'Он живёт в крупном городе, зарабатывает лично от $100 до $300 тыс. в год. Его бизнесу 3–8 лет. '
    'Он уже чего-то добился, но чувствует, что рынок этого не видит. Он хочет, чтобы его компания '
    'получила признание — не ради тщеславия, а потому что это работает: награда на сайте повышает '
    'доверие клиентов, помогает закрывать сделки, привлекать инвестиции и выделиться среди конкурентов. '
    'Он гуглит «business awards» не потому что ему скучно, а потому что увидел, что конкурент '
    'получил награду и теперь использует её в маркетинге.'
)

doc.add_heading('1.2. Демография — точные параметры', level=2)
add_table(
    ['Параметр', 'Значение', 'Примечание для таргетинга'],
    [
        ['Возраст', '33–48 лет (ядро), допустимо 28–55', 'Средний возраст успешного основателя в США — 42 года (MIT/Census). Топ-0.1% — 45 лет'],
        ['Пол', '58% мужчины / 42% женщины', '12.3 млн бизнесов в США принадлежат женщинам. Не исключать'],
        ['Семейное положение', 'Преимущественно женат/замужем', '65–70% владельцев бизнеса в этом возрасте в браке'],
        ['Образование', 'Бакалавр и выше', '~40% имеют MBA или магистратуру'],
        ['Личный доход', '$100,000–$350,000/год', 'Средняя зарплата CEO стартапа — $165,000 (2026)'],
        ['Доход домохозяйства', 'Top 10%–25% (от $150,000)', 'Для Google Ads: Household Income — Top 10% или Top 25%'],
        ['Выручка бизнеса', '$500,000–$10,000,000', '30.4% бизнесов в США имеют выручку > $1 млн'],
        ['Размер компании', '5–200 сотрудников', '89.8% бизнесов — менее 20 сотрудников, наш покупатель — подросший'],
        ['Стаж бизнеса', '3–8 лет', '62.3% бизнесов существуют 6+ лет'],
        ['Должность', 'CEO / Founder / Co-Founder / VP Marketing / Managing Director', 'Ключевые должности для таргетинга'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_heading('1.3. География — где он живёт', level=2)
add_para('Города первого приоритета (повышенные ставки +20–30%):', bold=True)
cities_1 = [
    'New York, NY', 'Los Angeles, CA', 'San Francisco / San Jose, CA (Silicon Valley)',
    'Miami, FL', 'Chicago, IL', 'Austin, TX', 'Boston, MA', 'Seattle, WA',
    'Dallas / Houston, TX'
]
for c in cities_1:
    add_bullet(c)

add_para('Города второго приоритета:', bold=True)
cities_2 = [
    'Atlanta, GA', 'Denver, CO', 'Phoenix, AZ', 'San Diego, CA', 'Washington, DC',
    'Portland, OR', 'Nashville, TN', 'Charlotte, NC', 'Raleigh-Durham, NC',
    'Minneapolis, MN', 'Las Vegas, NV'
]
for c in cities_2:
    add_bullet(c)

add_para('Для Google Ads: Таргетинг — все 50 штатов США, но с повышенными ставками (+20–30%) '
         'на города первого приоритета.', italic=True)

doc.add_heading('1.4. Отрасли — в каком бизнесе он работает', level=2)
add_table(
    ['Отрасль', 'Почему ищет награды', 'Подходящая площадка'],
    [
        ['Технологии / SaaS', 'Конкурентный рынок, нужно социальное доказательство для B2B-продаж', 'MYRONYX, Global 1000'],
        ['Финтех', 'Доверие критично, награда = сигнал для клиентов', 'Global 1000, MYRONYX'],
        ['Digital-маркетинг / агентства', 'Сами знают силу PR, ищут наградной контент', 'MYRONYX, NOVA'],
        ['Здоровье / Health-tech', 'Регулируемая отрасль, нужна внешняя валидация', 'MYRONYX, Velto'],
        ['Beauty / Wellness', 'Высокая конкуренция, награда = дифференциация', 'MYRONYX, Glonary, Velto'],
        ['Кибербезопасность', 'Доверие = основа бизнеса', 'Glonary'],
        ['E-commerce', 'Переполненный рынок, нужен «бейдж доверия» на сайте', 'MYRONYX, Glonary'],
        ['Проф. услуги (консалтинг, юр., фин.)', 'Личный бренд = продажи', 'Global 1000, NOVA, Velto'],
        ['AI / Machine Learning', 'Горячая ниша, все хотят выделиться', 'MYRONYX, NOVA, Velto'],
        ['Инфобизнес / коучинг', 'Нужно «доказать» экспертизу аудитории', 'Glonary'],
        ['Sustainability / Green-tech', 'ESG-повестка, награда подтверждает влияние', 'MYRONYX, Global 1000, Velto'],
    ],
    col_widths=[1.8, 2.5, 2.2]
)

doc.add_page_break()

# ============================================================
# PART 2
# ============================================================
doc.add_heading('Часть 2. Что он хочет получить', level=1)

doc.add_heading('2.1. Конечные цели — зачем ему награда', level=2)
add_para('Этот человек не покупает «трофей на полку». Он покупает бизнес-результат:', bold=True)

add_table(
    ['Цель', 'Что это значит конкретно', 'Важность'],
    [
        ['Доверие клиентов', 'Бейдж «Award Winner» на сайте → клиент выбирает его. 84% C-suite доверяют больше компании с наградой', 'Критично'],
        ['Закрытие крупных сделок', '«Мы — лауреаты [Award]» в презентации. 91% учитывают при выборе партнёра', 'Критично'],
        ['Привлечение инвестиций', 'Сигнал для инвестора о верификации. 89% учитывают при инвест. решениях', 'Высоко'],
        ['Контент для маркетинга', 'Пресс-релиз, пост в LinkedIn, баннер, строчка в email-подписи', 'Высоко'],
        ['SEO и бэклинки', 'Упоминание на сайте награды = бэклинк, пресс-релиз = цитирование', 'Средне-высоко'],
        ['Привлечение талантов', '76% кандидатов положительно реагируют на награды работодателя', 'Средне'],
        ['Личный бренд', 'LinkedIn-пост «Proud to be recognized...» → рост аудитории', 'Высоко'],
        ['Обоснование цен', '56% потребителей готовы платить премию за продукты лауреата', 'Средне'],
        ['Моральное удовлетворение', 'Подтверждение «я не самозванец» (85% основателей — imposter syndrome)', 'Высоко (скрытое)'],
    ],
    col_widths=[1.5, 3.5, 1.0]
)

doc.add_heading('2.2. Страхи и возражения', level=2)

add_table(
    ['Страх', 'Что он думает', 'Как снять в рекламе / на лендинге'],
    [
        ['«Это развод»', '«Очередная vanity award, которая ничего не стоит»', 'Реальные судьи с профилями, прозрачные критерии, известные прошлые победители'],
        ['«Никто не знает»', '«Мои клиенты не впечатлятся»', 'PR-покрытие: Business Insider, Yahoo Finance. Бейдж на сайтах реальных компаний'],
        ['«Слишком дорого»', '«$500 за заявку — а вдруг не выиграю?»', 'ROI-аргумент: «1 сделка окупает 100 заявок». Gold/Silver/Bronze — шанс выше'],
        ['«Я не подхожу»', '«Мой бизнес слишком маленький»', 'Кнопка «Check Eligibility», примеры победителей разного масштаба'],
        ['«Нет времени»', '«Анкета на 20 страниц»', '«Simple 3-part submission: 500 слов + background + links»'],
    ],
    col_widths=[1.2, 2.0, 3.3]
)

doc.add_page_break()

# ============================================================
# PART 3
# ============================================================
doc.add_heading('Часть 3. Как он ведёт себя в интернете', level=1)

doc.add_heading('3.1. Что он гуглит — поисковые запросы', level=2)

add_para('Этап 1: Начальный интерес (видел награду конкурента)', bold=True)
for kw in ['business awards 2026', 'best business awards to apply for',
           'business awards for small companies', 'innovation awards 2026', 'global business awards']:
    add_bullet(kw)

add_para('Этап 2: Поиск по своей отрасли', bold=True)
for kw in ['technology awards for startups', 'health and wellness awards 2026',
           'cybersecurity awards', 'beauty industry awards', 'AI awards 2026',
           'fintech awards', 'sustainability awards for companies',
           'women in business awards', 'young entrepreneur awards']:
    add_bullet(kw)

add_para('Этап 3: Сравнение и выбор', bold=True)
for kw in ['best business awards programs', 'how to apply for business awards',
           'business awards worth applying for', 'top international business awards']:
    add_bullet(kw)

add_para('Этап 4: Проверка легитимности (критически важный!)', bold=True)
for kw in ['[award name] reviews', '[award name] legitimate',
           '[award name] past winners', 'are paid business awards worth it']:
    add_bullet(kw)

add_para('Этап 5: Готов действовать', bold=True)
for kw in ['[award name] nomination', '[award name] apply',
           '[award name] entry fee', '[award name] deadline 2026']:
    add_bullet(kw)

doc.add_heading('3.2. Какие сайты посещает', level=2)
add_table(
    ['Категория', 'Сайты', 'Зачем'],
    [
        ['Бизнес-медиа', 'Forbes, Inc., Entrepreneur, FastCompany, BusinessInsider', 'Новости, тренды, списки (Inc. 5000)'],
        ['Tech-медиа', 'TechCrunch, Wired, VentureBeat', 'Если из tech-сектора'],
        ['LinkedIn', 'linkedin.com', 'Основная соцсеть, 2–5 публикаций/неделю'],
        ['Бизнес-образование', 'HBR.org, MIT Sloan Review', 'Стратегия, менеджмент'],
        ['Подкасты', 'Spotify, Apple Podcasts', 'How I Built This, My First Million'],
        ['Email-рассылки', 'Morning Brew, The Hustle, TLDR', 'Ежедневный бизнес-дайджест'],
        ['Конкурирующие награды', 'StevieAwards.com, TitanAwards.com, GRA, BestInBiz', 'Сравнение программ'],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_heading('3.3. Когда он онлайн', level=2)
add_table(
    ['Время (EST)', 'Что делает', 'Вероятность клика'],
    [
        ['6:00–8:00', 'Почта, новости, LinkedIn на телефоне', 'Средняя'],
        ['8:00–10:00', 'Начало рабочего дня', 'Низкая'],
        ['10:00–12:00', 'Перерывы, LinkedIn, гуглит', 'ВЫСОКАЯ'],
        ['12:00–14:00', 'Обед, скроллинг новостей', 'ВЫСОКАЯ'],
        ['14:00–16:00', 'Встречи, звонки', 'Низкая'],
        ['16:00–18:00', 'Завершение дня, LinkedIn', 'ВЫСОКАЯ'],
        ['20:00–23:00', 'Вечер, планирование', 'Средне-высокая'],
    ],
    col_widths=[1.3, 2.5, 2.7]
)

doc.add_heading('3.4. Устройства', level=2)
add_table(
    ['Устройство', 'Доля трафика', 'Поведение'],
    [
        ['Десктоп / Ноутбук', '55–60%', 'Заполняет заявки, читает длинные материалы. Конверсия выше → ставка +15%'],
        ['Смартфон (iPhone)', '35–40%', 'LinkedIn, быстрые поиски, первое знакомство'],
        ['Планшет (iPad)', '5–10%', 'Вечерний просмотр'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

doc.add_page_break()

# ============================================================
# PART 4
# ============================================================
doc.add_heading('Часть 4. Психология — почему он купит', level=1)

doc.add_heading('5 психологических триггеров покупки', level=2)

triggers = [
    ('Триггер 1: «Конкурент уже получил награду»',
     'Он скроллит LinkedIn и видит пост конкурента: «Thrilled to announce that [Company] has been '
     'recognized with the [Award]!» У поста 500 лайков. Он думает: «Чёрт, наш продукт лучше, '
     'но они теперь выглядят круче».',
     'Your competitors are getting recognized. Are you?',
     '[X] companies in your industry have already been nominated. Submissions close [date].'),

    ('Триггер 2: «Нужно соц. доказательство для большой сделки»',
     'Готовит презентацию для крупного клиента или инвестора. Слайд «О нас» пустой. '
     'Конкурент — с тремя бейджами.',
     'Close bigger deals with third-party recognition',
     '84% of executives trust award-winning companies more. Get recognized.'),

    ('Триггер 3: «Хочу, чтобы меня наконец заметили»',
     '5 лет строил бизнес. Вырос до $2 млн. Но рынок не знает. Синдром самозванца шепчет: '
     '«Может, ты не настолько хорош».',
     'Built something great? The world should know.',
     'Nominations now open. Recognition your business deserves.'),

    ('Триггер 4: «Это умная инвестиция, не трата»',
     'Считает деньги. $500 — не мало. Но: «Если выиграю — пресс-релиз, бейдж, LinkedIn-пост. '
     'Один новый клиент окупит в 10 раз».',
     '$500 investment → press coverage, trust badge, new clients',
     'Award winners report 2.8x more trust from prospects.'),

    ('Триггер 5: «Дедлайн приближается»',
     'FOMO + ограниченное время. «Final deadline: June 30, 2026.» '
     'Страх упустить сильнее страха потратить.',
     'Final nomination deadline: [Date]. Apply now.',
     'Only [X] spots remaining. Submit before [date].'),
]

for title, situation, headline, desc in triggers:
    doc.add_heading(title, level=3)
    add_para(f'Ситуация: {situation}', italic=True)
    add_para(f'Заголовок объявления: {headline}', bold=True)
    add_para(f'Текст объявления: {desc}')

doc.add_heading('Внутренний монолог покупателя', level=2)

monologue = [
    ('Видит рекламу:', '«Хм, ещё одна award-программа... Но подожди, у них есть категория для моей ниши. И сроки — как раз сейчас. Давай посмотрю.»'),
    ('На лендинге:', '«Сайт выглядит серьёзно. Что за категории? Кто выигрывал? Сколько стоит? Не так дорого, как Stevie. Кто судьи? Надо погуглить...»'),
    ('Гуглит название:', '«Есть ли статьи? Пресс-релизы? BBB? Если нахожу реальное — возвращаюсь на сайт.»'),
    ('Принимает решение:', '«Worst case — потеряю $500 и получу обратную связь. Best case — бейдж, PR, контент. Дешевле одного дня работы агентства.»'),
]
for stage, thought in monologue:
    add_para(stage, bold=True)
    add_para(thought, italic=True, color=(0x55, 0x55, 0x55))

doc.add_page_break()

# ============================================================
# PART 5
# ============================================================
doc.add_heading('Часть 5. Настройки Google Ads — конкретные параметры', level=1)

doc.add_heading('5.1. Структура кампаний', level=2)
add_table(
    ['Кампания', 'Площадка', 'Тип', 'Бюджет/день'],
    [
        ['Campaign 1', 'MYRONYX Global Awards', 'Search', '$30–50'],
        ['Campaign 2', 'Global 1000 Award', 'Search', '$20–40'],
        ['Campaign 3', 'Glonary Awards', 'Search', '$30–50'],
        ['Campaign 4', 'NOVA Awards', 'Search', '$15–30'],
        ['Campaign 5', 'Velto Awards', 'Search', '$15–30'],
        ['Campaign 6', 'General (все площадки)', 'Search', '$40–60'],
        ['Campaign 7', 'Ремаркетинг', 'Display', '$15–25'],
        ['Campaign 8', 'Discovery / Demand Gen', 'Discovery', '$20–30'],
    ],
    col_widths=[1.3, 2.0, 1.0, 1.2]
)
add_para('Общий начальный бюджет: $5,550–$9,450 / месяц', bold=True, size=12)

doc.add_heading('5.2. Ключевые слова', level=2)

add_para('Группа A: Высокоинтентные (ищет награду прямо сейчас)', bold=True, color=(0x1B, 0x3A, 0x5C))
kw_a = [
    '[business awards 2026]', '[business awards to apply for]', '[apply for business awards]',
    '[best business awards]', '[global business awards]', '[international business awards]',
    '[innovation awards 2026]', '[leadership awards]', '[business excellence awards]',
    '[entrepreneur awards]', '"business awards nominations"', '"business awards for startups"',
    '"business awards for small business"', '"CEO awards"', '"founder awards"'
]
for kw in kw_a:
    add_bullet(kw)

add_para('Группа B: По отрасли', bold=True, color=(0x1B, 0x3A, 0x5C))
kw_b = [
    '[technology awards 2026]', '[cybersecurity awards]', '[health and wellness awards]',
    '[beauty industry awards]', '[AI awards]', '[fintech awards]', '[sustainability awards]',
    '[digital marketing awards]', '[women in business awards]', '[young entrepreneur awards]',
    '[startup awards]', '[SaaS awards]', '"awards for beauty brands"', '"awards for healthcare companies"'
]
for kw in kw_b:
    add_bullet(kw)

add_para('Группа C: Конкурентные (перехват трафика)', bold=True, color=(0x1B, 0x3A, 0x5C))
kw_c = ['"stevie awards alternative"', '"titan awards alternative"', '"business awards like stevie"',
         '"globee awards"', '"global recognition awards"']
for kw in kw_c:
    add_bullet(kw)

add_para('Группа D: Брендовые', bold=True, color=(0x1B, 0x3A, 0x5C))
kw_d = ['[myronyx global awards]', '[global 1000 award]', '[nova awards business]',
         '[glonary awards]', '[velto awards]', '[velto european awards]']
for kw in kw_d:
    add_bullet(kw)

add_para('НЕГАТИВНЫЕ ключевые слова (обязательно!):', bold=True, color=(0xCC, 0x00, 0x00))
neg = ['free awards', 'scam', 'fake', 'vanity', 'trophy shop', 'custom trophies',
       'employee of the month', 'award certificate template', 'Grammy', 'Oscar', 'Emmy',
       'Nobel', 'music awards', 'film awards', 'school awards', 'kids awards', 'sports awards', 'military awards']
for kw in neg:
    add_bullet(f'-{kw}')

doc.add_heading('5.3. Аудитории (Audiences)', level=2)

add_para('In-Market аудитории (режим Observation → потом усилить ставки):', bold=True)
for a in ['/Business Services', '/Business Services/Advertising & Marketing Services',
          '/Business Services/Business Financial Services',
          '/Software/Business & Productivity Software', '/Business & Industrial Products']:
    add_bullet(a)

add_para('Affinity аудитории:', bold=True)
for a in ['Business Professionals', 'Technophiles', 'Avid Investors',
          'Frequently Attends Live Events', 'Avid News Readers']:
    add_bullet(a)

add_para('Custom Segment 1 — «Ищущие бизнес-награды»:', bold=True)
add_para('Ключевые слова: business awards, innovation awards, company recognition, '
         'business excellence, award nominations')

add_para('Custom Segment 2 — «Предприниматели»:', bold=True)
add_para('URL: forbes.com, inc.com, entrepreneur.com, techcrunch.com, fastcompany.com, hbr.org, linkedin.com')

add_para('Custom Segment 3 — «Конкуренты»:', bold=True)
add_para('URL: stevieawards.com, titanawards.com, globalrecognitionawards.org, bestinbizawards.com, globeeawards.com')

add_para('Detailed Demographics:', bold=True)
for d in ['Education: Bachelor\'s degree or higher',
          'Household Income: Top 10% or Top 25%',
          'Company Size: Small employer (1–249) or Midsize (250–999)',
          'Homeownership: Homeowner']:
    add_bullet(d)

doc.add_heading('5.4. Расписание показов', level=2)
add_table(
    ['День', 'Время (EST)', 'Ставка'],
    [
        ['Пн–Пт', '6:00–9:00', 'Стандартная (100%)'],
        ['Пн–Пт', '9:00–12:00', 'Повышенная (+25%)'],
        ['Пн–Пт', '12:00–14:00', 'Повышенная (+20%)'],
        ['Пн–Пт', '14:00–16:00', 'Стандартная (100%)'],
        ['Пн–Пт', '16:00–19:00', 'Повышенная (+20%)'],
        ['Пн–Пт', '19:00–23:00', 'Стандартная (100%)'],
        ['Суббота', '9:00–14:00', 'Стандартная (100%)'],
        ['Воскресенье', '—', 'Выкл или пониженная (-30%)'],
    ],
    col_widths=[1.3, 1.5, 3.7]
)

doc.add_heading('5.5. Устройства', level=2)
add_table(
    ['Устройство', 'Коррекция ставки'],
    [
        ['Десктоп / Ноутбук', '+15% (выше конверсия)'],
        ['Мобильный', '0% (первый контакт часто с телефона)'],
        ['Планшет', '0%'],
    ],
    col_widths=[2.5, 4.0]
)

doc.add_page_break()

# ============================================================
# PART 6
# ============================================================
doc.add_heading('Часть 6. Примеры объявлений', level=1)

ads = [
    ('Общая кампания — объявление 1 (конкурентный триггер)',
     ['Your Competitors Are Getting Recognized', 'Global Business Awards 2026', 'Nominations Now Open'],
     '84% of executives trust award-winning companies more. Get the recognition your company deserves. Apply in 10 minutes.',
     'Join CEOs, founders & innovators recognized globally. Simple 3-step application. Multiple categories. Expert judging panel.'),

    ('Общая кампания — объявление 2 (ROI-триггер)',
     ['Turn $500 Into Press Coverage & Trust', 'Business Awards — Apply Now', 'Gold · Silver · Bronze Recognition'],
     'Award winners report 2.8x higher trust from clients. Boost credibility with independent third-party recognition. Apply today.',
     'Get a trust badge for your website, ready-made press release, and LinkedIn-worthy recognition. Expert-judged. Transparent criteria.'),

    ('Общая кампания — объявление 3 (дедлайн/FOMO)',
     ['Nomination Deadline Approaching', '2026 Business Excellence Awards', "Don't Miss Your Chance"],
     'Submissions close soon. Nominate your company for global recognition in business, tech, health, innovation & more.',
     'Limited spots per category. Multi-stage expert review. Past winners include top companies worldwide. Check eligibility now.'),

    ('MYRONYX Global Awards',
     ['MYRONYX Global Awards 2026', 'Business · Digital · Health Excellence', 'Ceremony in New York'],
     'Global recognition for builders of tomorrow. 9+ categories: Innovation, Leadership, Technology, Health & Wellness. Apply now.',
     'Expert judging panel. Gold, Silver & Bronze awards. Get press coverage, trust badge & global recognition. NYC ceremony.'),

    ('Global 1000 Award',
     ['Global 1000 Award — Exclusive', '1,000 Visionaries. One Standard.', 'By Invitation & Nomination Only'],
     'The most selective global business recognition. Limited to 1,000 exceptional leaders. Blind-reviewed.',
     'Are you a Visionary CEO, Tech Disruptor or Social Pioneer? Join the Founders\' Circle. Exclusive merit-based recognition.'),

    ('Glonary Awards',
     ['Glonary Awards — Global Recognition', 'CyberSecurity · Women · Innovation', 'Ceremony in Las Vegas'],
     '9 award programs: Young Innovators, Women in Business, Technology & Design, Business Leaders. Apply for 2026.',
     'Peer-to-peer evaluation. Las Vegas ceremony. Global recognition for leaders, innovators and changemakers.'),

    ('NOVA Awards',
     ['NOVA Awards — Proof Over Hype', 'Where Industry Futures Are Recognized', 'Multi-Stage Expert Review'],
     'Selective European awards for visionaries, innovators & digital leaders. Multi-stage advisory review.',
     'Categories: Digital Security, Leadership, Innovation, Emerging Talent. Rigorous expert judging. Apply now.'),

    ('Velto European Awards',
     ['Velto Awards — Elite Recognition', 'European Business Excellence', 'Gold · Silver · Bronze'],
     'Rigorous 10-point scoring by 50+ volunteer judges. Categories: Business Growth, AI, Beauty & Wellness, Sustainability.',
     'Elite recognition for industry leaders and innovators. Multi-stage evaluation. Digital-first. Transparent criteria.'),
]

for title, headlines, desc1, desc2 in ads:
    doc.add_heading(title, level=3)
    for i, h in enumerate(headlines):
        add_para(f'Headline {i+1}: {h}', bold=True)
    add_para(f'Description 1: {desc1}')
    add_para(f'Description 2: {desc2}')

doc.add_page_break()

# ============================================================
# PART 7
# ============================================================
doc.add_heading('Часть 7. Путь клиента от рекламы до оплаты', level=1)

add_para('Среднее время от первого клика до оплаты: 7–21 день', bold=True, size=13,
         color=(0xCC, 0x00, 0x00))

stages = [
    ('ДЕНЬ 1 — Первый контакт',
     ['Видит рекламу в Google (Search или Display)',
      'Кликает → попадает на лендинг площадки',
      'Просматривает категории, судей, прошлых победителей',
      'УХОДИТ (80% уходят с первого визита)']),
    ('ДЕНЬ 2–5 — Проверка',
     ['Гуглит название площадки → проверяет легитимность',
      'Читает прошлых победителей',
      'Ищет отзывы, пресс-релизы, BBB',
      'Видит ремаркетинг-рекламу (Display)']),
    ('ДЕНЬ 5–14 — Решение',
     ['Возвращается на сайт через ремаркетинг или прямой заход',
      'Изучает конкретную категорию, стоимость, процесс',
      'Сравнивает с 1–2 другими программами наград',
      'Обсуждает с партнёром/маркетологом/ассистентом']),
    ('ДЕНЬ 7–21 — Действие',
     ['Принимает решение подать заявку',
      'Заполняет форму (10–30 минут)',
      'Оплачивает entry fee',
      '→ ЛИД → ПРОДАЖА']),
    ('ПОСЛЕ ПОБЕДЫ — Виральный цикл',
     ['Получает результат',
      'Публикует в LinkedIn: «Proud to announce...»',
      'Ставит бейдж на сайт',
      'Его пост видят конкуренты → гуглят → ЦИКЛ ПОВТОРЯЕТСЯ']),
]

for title, items in stages:
    doc.add_heading(title, level=3)
    for item in items:
        add_bullet(item)

add_para('')
add_para('ВЫВОД: Ремаркетинг критически важен. Минимум 70% конверсий — не с первого клика.',
         bold=True, size=12, color=(0xCC, 0x00, 0x00))

doc.add_page_break()

# ============================================================
# PART 8
# ============================================================
doc.add_heading('Часть 8. Метрики и KPI', level=1)

doc.add_heading('8.1. Целевые показатели', level=2)
add_table(
    ['Метрика', 'Целевое значение', 'Как считать'],
    [
        ['CPC (стоимость клика)', '$2.50–$5.00', 'Средняя по кампании'],
        ['CTR (кликабельность)', '4–8%', 'Клики / Показы'],
        ['Конверсия (заявка)', '3–7%', 'Заявки / Клики'],
        ['CPA (стоимость заявки)', '$40–$80', 'Бюджет / Заявки'],
        ['ROAS', '300–500%', 'Доход / Расход на рекламу'],
        ['Средний чек', '$400–$700', 'Entry fee'],
        ['LTV (3 года)', '$1,500–$4,000', 'Повторные заявки + upsells'],
    ],
    col_widths=[2.0, 1.5, 3.0]
)

doc.add_heading('8.2. Точки конверсии', level=2)
add_bullet('Основная: ', bold_prefix='')
add_para('Оплаченная заявка (entry fee paid)', bold=True)
add_bullet('Мягкая конверсия 1: Заполненная форма заявки (без оплаты)')
add_bullet('Мягкая конверсия 2: Клик по «Check Eligibility» / «Start Application»')
add_bullet('Мягкая конверсия 3: Время на сайте > 3 мин + просмотр > 3 страниц')

doc.add_page_break()

# ============================================================
# PART 9
# ============================================================
doc.add_heading('Часть 9. Резюме — кто этот американец', level=1)

add_para(
    'Это мужчина или женщина 35–45 лет, из Нью-Йорка, Лос-Анджелеса, Майами, Сан-Франциско '
    'или другого крупного города. Основатель или CEO компании с 10–100 сотрудниками и выручкой '
    '$1–5 млн. Высшее образование, доход домохозяйства в топ-25% ($150K+). Работает в технологиях, '
    'финтехе, здоровье, beauty или digital-сервисах. Активен на LinkedIn.',
    size=12
)
add_para(
    'Его бизнесу 3–7 лет. Он уже успешен, но рынок об этом не знает. Он ищет «business awards 2026» '
    'в Google, потому что увидел, что конкурент получил награду и использует её в маркетинге.',
    size=12
)
add_para(
    'Он хочет получить бейдж на сайт, пресс-релиз, пост в LinkedIn и аргумент для инвесторов/клиентов. '
    'Он готов заплатить $400–$700 за заявку, если убедится, что награда настоящая.',
    size=12
)
add_para(
    'Его главный страх — что это «fake award». Его главная мотивация — социальное доказательство, '
    'закрытие сделок и признание.',
    size=12
)
add_para(
    'Время от первого клика до оплаты — 7–21 день. Ремаркетинг обязателен.',
    size=12, bold=True
)

doc.add_paragraph()
doc.add_paragraph()
add_para('Источники: MIT, Census Bureau, SBA, Constant Contact, GRA Credibility Effect 2026, '
         'Awards Trust Mark, BBB, WordStream, Kruze Consulting, отраслевые бенчмарки 2025–2026.',
         size=9, italic=True, color=(0x99, 0x99, 0x99))

# ============================================================
# SAVE
# ============================================================
doc.save('/workspace/GOOGLE_ADS_BUYER_PERSONA.docx')
print('DONE: /workspace/GOOGLE_ADS_BUYER_PERSONA.docx')
